import os
import re
from datetime import datetime
from pathlib import Path
from copy import deepcopy
from docx import Document
from docx.shared import Pt, Inches

class ValidadorDados:
    @staticmethod
    def validar_cpf(cpf):
        """Remove caracteres especiais e valida estrutura básica"""
        cpf = re.sub(r'\D', '', cpf)
        if len(cpf) != 11:
            raise ValueError(f"CPF inválido: deve ter 11 dígitos. Recebido: {cpf}")
        return cpf
    
    @staticmethod
    def validar_cnpj(cnpj):
        """Remove caracteres especiais e valida estrutura básica"""
        cnpj = re.sub(r'\D', '', cnpj)
        if len(cnpj) != 14:
            raise ValueError(f"CNPJ inválido: deve ter 14 dígitos. Recebido: {cnpj}")
        return cnpj
    
    @staticmethod
    def formatar_cpf(cpf):
        """Formata CPF: 12345678901 → 123.456.789-01"""
        cpf = re.sub(r'\D', '', cpf)
        return f"{cpf[:3]}.{cpf[3:6]}.{cpf[6:9]}-{cpf[9:]}"
    
    @staticmethod
    def formatar_cnpj(cnpj):
        """Formata CNPJ: 12345678000190 → 12.345.678/0001-90"""
        cnpj = re.sub(r'\D', '', cnpj)
        return f"{cnpj[:2]}.{cnpj[2:5]}.{cnpj[5:8]}/{cnpj[8:12]}-{cnpj[12:]}"


class Gerador:
    TIPOS_LICENCA = {
        '1': 'Básica',
        '2': 'Profissional',
        '3': 'Enterprise',
        '4': 'Premium',
        '5': 'Startup',
        '6': 'Corporativa'
    }
    
    def __init__(self):
        self.caminho_templates = Path('templates')
        self.caminho_saida = Path('contratos_gerados')
        self._criar_diretorios()
    
    def _criar_diretorios(self):
        """Cria diretórios se não existirem"""
        self.caminho_templates.mkdir(exist_ok=True)
        self.caminho_saida.mkdir(exist_ok=True)
    
    def _mapear_template(self, tipo_licenca):
        """Mapeia o tipo de licença ao arquivo de template"""
        nome_arquivo = f"template_licenca_{tipo_licenca.lower()}.docx"
        caminho_completo = self.caminho_templates / nome_arquivo
        
        if not caminho_completo.exists():
            raise FileNotFoundError(f"Template não encontrado: {caminho_completo}")
        
        return caminho_completo
    
    def _substituir_em_paragrafo(self, paragrafo, substituicoes):
        """Substitui variáveis em um parágrafo mantendo formatação"""
        for chave, valor in substituicoes.items():
            if chave in paragrafo.text:
                # Se a variável é um parágrafo inteiro, substitui direto
                if paragrafo.text == chave:
                    paragrafo.text = valor
                # Se é parte do parágrafo, substitui mantendo runs
                else:
                    for run in paragrafo.runs:
                        if chave in run.text:
                            run.text = run.text.replace(chave, valor)
    
    def _substituir_em_tabelas(self, doc, substituicoes):
        """Substitui variáveis em todas as tabelas do documento"""
        for tabela in doc.tables:
            for linha in tabela.rows:
                for celula in linha.cells:
                    for paragrafo in celula.paragraphs:
                        self._substituir_em_paragrafo(paragrafo, substituicoes)
    
    def _substituir_em_corpo(self, doc, substituicoes):
        """Substitui variáveis em todos os parágrafos do documento"""
        for paragrafo in doc.paragraphs:
            self._substituir_em_paragrafo(paragrafo, substituicoes)
    
    def gerar_contrato_pf(self, nome, cpf, tipo_licenca):
        """Gera contrato para pessoa física"""
        # Validação
        cpf_limpo = ValidadorDados.validar_cpf(cpf)
        cpf_formatado = ValidadorDados.formatar_cpf(cpf_limpo)
        
        # Carrega template
        caminho_template = self._mapear_template(self.TIPOS_LICENCA[tipo_licenca])
        doc = Document(caminho_template)
        
        # Prepara substituições
        substituicoes = {
            '{NOME_CLIENTE}': nome,
            '{CPF}': cpf_formatado,
            '{TIPO_LICENCA}': self.TIPOS_LICENCA[tipo_licenca],
            '{DATA}': datetime.now().strftime('%d de %B de %Y'),
            '{DATA_BR}': datetime.now().strftime('%d/%m/%Y'),
            '{QTD_EQUIPOS}': '1',  # Exemplo fixo, pode ser parametrizado
        }
        
        # Substitui em corpo e tabelas
        self._substituir_em_corpo(doc, substituicoes)
        self._substituir_em_tabelas(doc, substituicoes)
        
        # Salva arquivo
        nome_arquivo = f"{nome}_{cpf_limpo}_{self.TIPOS_LICENCA[tipo_licenca]}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        caminho_saida = self.caminho_saida / nome_arquivo
        doc.save(caminho_saida)
        
        return caminho_saida
    
    def gerar_contrato_pj(self, razao_social, cnpj, tipo_licenca):
        """Gera contrato para pessoa jurídica"""
        # Validação
        cnpj_limpo = ValidadorDados.validar_cnpj(cnpj)
        cnpj_formatado = ValidadorDados.formatar_cnpj(cnpj_limpo)
        
        # Carrega template
        caminho_template = self._mapear_template(self.TIPOS_LICENCA[tipo_licenca])
        doc = Document(caminho_template)
        
        # Prepara substituições
        substituicoes = {
            '{RAZAO_SOCIAL}': razao_social,
            '{CNPJ}': cnpj_formatado,
            '{TIPO_LICENCA}': self.TIPOS_LICENCA[tipo_licanca],
            '{DATA}': datetime.now().strftime('%d de %B de %Y'),
            '{DATA_BR}': datetime.now().strftime('%d/%m/%Y'),
        }
        
        # Substitui em corpo e tabelas
        self._substituir_em_corpo(doc, substituicoes)
        self._substituir_em_tabelas(doc, substituicoes)
        
        # Salva arquivo
        nome_arquivo = f"{razao_social}_{cnpj_limpo}_{self.TIPOS_LICENCA[tipo_licenca]}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        caminho_saida = self.caminho_saida / nome_arquivo
        doc.save(caminho_saida)
        
        return caminho_saida
    
    def listar_tipos_licenca(self):
        """Exibe opções de licença disponíveis"""
        print("\n--- TIPOS DE LICENÇA ---")
        for chave, nome in self.TIPOS_LICENCA.items():
            print(f"{chave} - {nome}")


def main():
    gerador = Gerador()
    
    print("\n=== GERADOR DE CONTRATOS ===\n")
    
    # Tipo de pessoa
    tipo_pessoa = input("Tipo de cliente (1=Pessoa Física, 2=Pessoa Jurídica): ").strip()
    
    if tipo_pessoa == '1':
        # Pessoa Física
        nome = input("Nome completo: ").strip()
        cpf = input("CPF (apenas números ou formatado): ").strip()
        
        gerador.listar_tipos_licenca()
        tipo_licenca = input("\nEscolha o tipo de licença (1-6): ").strip()
        
        try:
            caminho = gerador.gerar_contrato_pf(nome, cpf, tipo_licenca)
            print(f"\n✓ Contrato gerado com sucesso: {caminho}")
        except Exception as e:
            print(f"\n✗ Erro: {e}")
    
    elif tipo_pessoa == '2':
        # Pessoa Jurídica
        razao_social = input("Razão Social: ").strip()
        cnpj = input("CNPJ (apenas números ou formatado): ").strip()
        
        gerador.listar_tipos_licenca()
        tipo_licenca = input("\nEscolha o tipo de licença (1-6): ").strip()
        
        try:
            caminho = gerador.gerar_contrato_pj(razao_social, cnpj, tipo_licenca)
            print(f"\n✓ Contrato gerado com sucesso: {caminho}")
        except Exception as e:
            print(f"\n✗ Erro: {e}")
    
    else:
        print("Opção inválida.")


if __name__ == '__main__':
    main()